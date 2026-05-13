"""Export a logbook entry as a styled .docx file.

Parses the raw markdown body and builds a Word document with:
- Inter / JetBrains Mono fonts matching the web UI
- Styled headings (H1–H3), paragraphs, lists, blockquotes, code blocks
- Markdown tables rendered as Word tables with header shading
- Embedded images (fetched from the local image store)
- Figure captions styled as italic gray text
- Bold/italic/code/strikethrough/links inline formatting
- Entry metadata (project, ID, dates) in a subtle header
"""

import io
import os
import re

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .logbooks import get_image_path

# ── Palette (matches the web UI light theme) ──────────────────────────────────

_TEXT = RGBColor(0x18, 0x18, 0x1F)
_MUTED = RGBColor(0x90, 0x90, 0xAA)
_ACCENT = RGBColor(0x2B, 0xA2, 0x98)
_CODE_BG = RGBColor(0xF5, 0xF5, 0xF7)
_BORDER = RGBColor(0xE4, 0xE4, 0xEC)
_SURFACE = RGBColor(0xF5, 0xF5, 0xF7)
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

_FONT_SANS = "Inter"
_FONT_MONO = "JetBrains Mono"


# ── Style helpers ─────────────────────────────────────────────────────────────

def _setup_styles(doc):
    """Configure document-level styles to mirror the logbook web UI."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = _FONT_SANS
    font.size = Pt(11)
    font.color.rgb = _TEXT
    pf = style.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.5

    for level, size, before, after in [
        (1, 22, 24, 6),
        (2, 16, 20, 4),
        (3, 13, 14, 3),
    ]:
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = _FONT_SANS
        hs.font.size = Pt(size)
        hs.font.bold = True
        hs.font.color.rgb = _TEXT
        hs.paragraph_format.space_before = Pt(before)
        hs.paragraph_format.space_after = Pt(after)

    if "CodeBlock" not in [s.name for s in doc.styles]:
        cb = doc.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        cb.font.name = _FONT_MONO
        cb.font.size = Pt(9)
        cb.font.color.rgb = _TEXT
        cb.paragraph_format.space_before = Pt(6)
        cb.paragraph_format.space_after = Pt(6)
        _set_paragraph_shading(cb.paragraph_format, _CODE_BG)

    if "FigCaption" not in [s.name for s in doc.styles]:
        fc = doc.styles.add_style("FigCaption", WD_STYLE_TYPE.PARAGRAPH)
        fc.font.name = _FONT_SANS
        fc.font.size = Pt(9)
        fc.font.italic = True
        fc.font.color.rgb = _MUTED
        fc.paragraph_format.space_before = Pt(2)
        fc.paragraph_format.space_after = Pt(10)
        fc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if "BlockQuote" not in [s.name for s in doc.styles]:
        bq = doc.styles.add_style("BlockQuote", WD_STYLE_TYPE.PARAGRAPH)
        bq.font.name = _FONT_SANS
        bq.font.size = Pt(10)
        bq.font.italic = True
        bq.font.color.rgb = _MUTED
        bq.paragraph_format.space_before = Pt(4)
        bq.paragraph_format.space_after = Pt(4)
        bq.paragraph_format.left_indent = Pt(18)
        _set_left_border(bq, _ACCENT, Pt(2))

    if "ListBullet1" not in [s.name for s in doc.styles]:
        lb = doc.styles.add_style("ListBullet1", WD_STYLE_TYPE.PARAGRAPH)
        lb.font.name = _FONT_SANS
        lb.font.size = Pt(11)
        lb.font.color.rgb = _TEXT
        lb.paragraph_format.left_indent = Pt(24)
        lb.paragraph_format.space_before = Pt(1)
        lb.paragraph_format.space_after = Pt(1)

    if "EntryMeta" not in [s.name for s in doc.styles]:
        em = doc.styles.add_style("EntryMeta", WD_STYLE_TYPE.PARAGRAPH)
        em.font.name = _FONT_SANS
        em.font.size = Pt(9)
        em.font.color.rgb = _MUTED
        em.paragraph_format.space_before = Pt(0)
        em.paragraph_format.space_after = Pt(16)


def _set_paragraph_shading(pf, color):
    """Apply background shading to a paragraph format."""
    el = pf._element
    shading = el.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): f"{color}",
    })
    pPr = el if el.tag.endswith("}pPr") else el.find(qn("w:pPr"))
    if pPr is None:
        pPr = el.makeelement(qn("w:pPr"), {})
        el.insert(0, pPr)
    pPr.append(shading)


def _set_left_border(style, color, width):
    """Set a left border on a style (for blockquote styling)."""
    pPr = style._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = style._element.makeelement(qn("w:pPr"), {})
        style._element.insert(0, pPr)
    borders = pPr.makeelement(qn("w:pBdr"), {})
    left = borders.makeelement(qn("w:left"), {
        qn("w:val"): "single",
        qn("w:sz"): str(int(width.pt * 8)),
        qn("w:space"): "8",
        qn("w:color"): f"{color}",
    })
    borders.append(left)
    pPr.append(borders)


def _shade_cell(cell, color):
    """Set background fill on a table cell."""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = tc.makeelement(qn("w:tcPr"), {})
        tc.insert(0, tcPr)
    shading = tcPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): f"{color}",
    })
    tcPr.append(shading)


def _set_cell_borders(cell, color_hex="E4E4EC"):
    """Set thin borders on a table cell."""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = tc.makeelement(qn("w:tcPr"), {})
        tc.insert(0, tcPr)
    borders = tcPr.makeelement(qn("w:tcBorders"), {})
    for side in ("top", "bottom", "left", "right"):
        el = borders.makeelement(qn(f"w:{side}"), {
            qn("w:val"): "single",
            qn("w:sz"): "4",
            qn("w:space"): "0",
            qn("w:color"): color_hex,
        })
        borders.append(el)
    tcPr.append(borders)


# ── Inline markdown parsing ──────────────────────────────────────────────────

_INLINE_PATTERNS = [
    ("bold", re.compile(r"\*\*(.+?)\*\*")),
    ("latex_bold", re.compile(r"\\textbf\{([^{}]+)\}")),
    ("small_caps", re.compile(r"\\textsc\{([^{}]+)\}")),
    ("italic", re.compile(r"\*(.+?)\*")),
    ("code", re.compile(r"`(.+?)`")),
    ("strike", re.compile(r"~~(.+?)~~")),
    ("link", re.compile(r"\[([^\]]+)\]\(((?:https?://|/|\.{0,2}/|#)[^)]+)\)")),
    ("image", re.compile(r"!\[([^\]]*)\]\(((?:https?://|/|\.{0,2}/)[^)]+)\)")),
]


def _add_inline_runs(paragraph, text, project=None, default_bold=False):
    """Parse inline markdown (bold, italic, code, links) and add runs."""
    if not text:
        return

    segments = _split_inline(text)
    for kind, content, extra in segments:
        run = paragraph.add_run(content)
        run.font.name = _FONT_SANS
        if default_bold:
            run.font.bold = True
        if kind in ("bold", "latex_bold"):
            run.font.bold = True
        elif kind == "small_caps":
            run.font.small_caps = True
            run.font.bold = True
        elif kind == "italic":
            run.font.italic = True
        elif kind == "code":
            run.font.name = _FONT_MONO
            run.font.size = Pt(9)
            _shade_run(run, _CODE_BG)
        elif kind == "strike":
            run.font.strike = True
        elif kind == "link":
            run.font.color.rgb = _ACCENT
            run.font.underline = True


def _split_inline(text):
    """Split text into segments of (kind, content, extra)."""
    segments = []
    pos = 0

    combined = re.compile(
        r"(?P<image>!\[(?P<ialt>[^\]]*)\]\((?P<isrc>(?:https?://|/|\.{0,2}/)[^)]+)\))"
        r"|(?P<latex_bold>\\textbf\{(?P<lbcont>[^{}]+)\})"
        r"|(?P<small_caps>\\textsc\{(?P<sccont>[^{}]+)\})"
        r"|(?P<bold>\*\*(?P<bcont>.+?)\*\*)"
        r"|(?P<italic>\*(?P<icont>.+?)\*)"
        r"|(?P<code>`(?P<ccont>.+?)`)"
        r"|(?P<strike>~~(?P<scont>.+?)~~)"
        r"|(?P<link>\[(?P<ltxt>[^\]]+)\]\((?P<lurl>(?:https?://|/|\.{0,2}/|#)[^)]+)\))"
    )

    for m in combined.finditer(text):
        if m.start() > pos:
            segments.append(("text", text[pos:m.start()], None))

        if m.group("image"):
            segments.append(("image", m.group("ialt"), m.group("isrc")))
        elif m.group("latex_bold"):
            segments.append(("latex_bold", m.group("lbcont"), None))
        elif m.group("small_caps"):
            segments.append(("small_caps", m.group("sccont"), None))
        elif m.group("bold"):
            segments.append(("bold", m.group("bcont"), None))
        elif m.group("italic"):
            segments.append(("italic", m.group("icont"), None))
        elif m.group("code"):
            segments.append(("code", m.group("ccont"), None))
        elif m.group("strike"):
            segments.append(("strike", m.group("scont"), None))
        elif m.group("link"):
            segments.append(("link", m.group("ltxt"), m.group("lurl")))

        pos = m.end()

    if pos < len(text):
        segments.append(("text", text[pos:], None))

    return segments


def _shade_run(run, color):
    """Apply highlight/shading to an inline run."""
    rPr = run._r.find(qn("w:rPr"))
    if rPr is None:
        rPr = run._r.makeelement(qn("w:rPr"), {})
        run._r.insert(0, rPr)
    shading = rPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): f"{color}",
    })
    rPr.append(shading)


# ── Table row helpers ─────────────────────────────────────────────────────────

def _is_table_row(line):
    t = line.strip()
    return t.startswith("|") and t.endswith("|") and "|" in t[1:-1]


def _is_table_sep(line):
    return bool(re.match(r"^\|[\s:|\-]+\|$", line.strip()))


def _parse_table_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


# ── Figure caption detection ─────────────────────────────────────────────────

_FIGURE_CAPTION_RE = re.compile(r"^\*{0,3}\*?Figure\s+\d+", re.IGNORECASE)


def _is_figure_caption(text):
    return bool(_FIGURE_CAPTION_RE.match(text.strip()))


# ── Image embedding ──────────────────────────────────────────────────────────

def _try_add_image(doc, project, src, width=Inches(5.5)):
    """Attempt to embed an image from the logbook image store."""
    m = re.match(r"/api/logbook/([^/]+)/images/(.+)$", src)
    if not m:
        return False

    img_project = m.group(1)
    filename = m.group(2)
    path = get_image_path(img_project, filename)
    if not path:
        return False

    ext = os.path.splitext(path)[1].lower()
    if ext in (".html", ".htm"):
        p = doc.add_paragraph()
        run = p.add_run(f"[Interactive figure: {filename}]")
        run.font.color.rgb = _MUTED
        run.font.italic = True
        run.font.size = Pt(10)
        return True

    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=width)
        return True
    except Exception:
        return False


# ── Main export function ──────────────────────────────────────────────────────

def export_entry_docx(project, entry):
    """Build a .docx from a logbook entry dict and return bytes."""
    doc = Document()

    _set_narrow_margins(doc)
    _setup_styles(doc)

    # Title
    title_p = doc.add_heading(entry.get("title", "Untitled"), level=1)
    title_p.runs[0].font.size = Pt(24)

    # Metadata line
    meta_parts = [project, f"#{entry.get('id', '')}"]
    if entry.get("created_at"):
        meta_parts.append(f"Created {entry['created_at']}")
    if entry.get("edited_at") and entry["edited_at"] != entry.get("created_at"):
        meta_parts.append(f"Edited {entry['edited_at']}")
    if entry.get("entry_type") == "plan":
        meta_parts.append("PLAN")
    elif entry.get("entry_type") == "campaign_board":
        meta_parts.append("CAMPAIGN BOARD")
        if entry.get("campaign"):
            meta_parts.append(f"campaign:{entry['campaign']}")
    meta_p = doc.add_paragraph(style="EntryMeta")
    meta_p.add_run(" · ".join(meta_parts))

    cg = (entry.get("campaign_goal") or "").strip()
    if entry.get("entry_type") == "campaign_board" and cg:
        doc.add_paragraph(style="Body")
        doc.add_paragraph(cg, style="Body")

    body = entry.get("body", "")
    _render_body(doc, body, project)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _set_narrow_margins(doc):
    """Set page margins to 1 inch all around."""
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def _render_body(doc, body, project):
    """Parse markdown body line by line and add elements to the document."""
    lines = body.split("\n")
    in_code = False
    code_lines = []
    table_buffer = []
    last_was_figure = False
    caption_buffer = []

    def flush_code():
        nonlocal code_lines
        if code_lines:
            text = "\n".join(code_lines)
            p = doc.add_paragraph(style="CodeBlock")
            run = p.add_run(text)
            run.font.name = _FONT_MONO
            run.font.size = Pt(9)
            code_lines = []

    def flush_table():
        nonlocal table_buffer
        if not table_buffer:
            return
        rows = table_buffer
        table_buffer = []

        header_cells = _parse_table_row(rows[0])
        data_start = 2 if len(rows) > 1 and _is_table_sep(rows[1]) else 1
        data_rows = [_parse_table_row(r) for r in rows[data_start:]]

        ncols = len(header_cells)
        nrows = 1 + len(data_rows)

        tbl = doc.add_table(rows=nrows, cols=ncols)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = True

        for ci, text in enumerate(header_cells):
            cell = tbl.rows[0].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            _add_inline_runs(p, text, project, default_bold=True)
            p.runs[0].font.size = Pt(9) if p.runs else None
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.bold = True
                run.font.name = _FONT_SANS
            _shade_cell(cell, _SURFACE)
            _set_cell_borders(cell)

        for ri, row_cells in enumerate(data_rows):
            for ci, text in enumerate(row_cells):
                if ci >= ncols:
                    break
                cell = tbl.rows[ri + 1].cells[ci]
                cell.text = ""
                p = cell.paragraphs[0]
                _add_inline_runs(p, text, project)
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = _FONT_SANS
                _set_cell_borders(cell)

    def flush_caption():
        nonlocal last_was_figure, caption_buffer
        if not last_was_figure:
            return
        if caption_buffer:
            text = " ".join(caption_buffer)
            p = doc.add_paragraph(style="FigCaption")
            _add_inline_runs(p, text, project)
            caption_buffer = []
        last_was_figure = False

    i = 0
    while i < len(lines):
        line = lines[i]

        # Code fence
        if line.strip().startswith("```"):
            flush_caption()
            flush_table()
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                flush_code()
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Blockquote / figure caption
        qm = re.match(r"^>\s?(.*)$", line)
        if qm:
            content = qm.group(1)
            if last_was_figure and (caption_buffer or _is_figure_caption(content)):
                caption_buffer.append(content)
                i += 1
                continue
            flush_caption()
            flush_table()
            p = doc.add_paragraph(style="BlockQuote")
            _add_inline_runs(p, content, project)
            i += 1
            continue

        # Empty line
        if not line.strip():
            if not last_was_figure:
                pass  # just skip blank lines
            i += 1
            continue

        flush_caption()

        # Table rows
        if _is_table_row(line):
            table_buffer.append(line)
            i += 1
            continue
        flush_table()

        # Headings
        hm = re.match(r"^(#{1,3})\s+(.*)$", line)
        if hm:
            level = len(hm.group(1))
            doc.add_heading(hm.group(2), level=level)
            i += 1
            continue

        # List items
        lm = re.match(r"^\s*[-*]\s+(.*)$", line)
        if lm:
            p = doc.add_paragraph(style="ListBullet1")
            run = p.add_run("• ")
            run.font.name = _FONT_SANS
            _add_inline_runs(p, lm.group(1), project)
            i += 1
            continue

        # Standalone image line
        img_match = re.match(
            r"^\s*!\[([^\]]*)\]\(((?:https?://|/|\.{0,2}/)[^)]+)\)\s*$", line
        )
        if img_match:
            flush_table()
            src = img_match.group(2)
            if not _is_html_embed(src):
                if _try_add_image(doc, project, src):
                    last_was_figure = True
                else:
                    p = doc.add_paragraph()
                    run = p.add_run(f"[Image: {src}]")
                    run.font.color.rgb = _MUTED
                    run.font.italic = True
            else:
                p = doc.add_paragraph()
                run = p.add_run(f"[Interactive figure: {src}]")
                run.font.color.rgb = _MUTED
                run.font.italic = True
            i += 1
            continue

        # HTML embed (bare URL on its own line)
        if _is_html_embed(line.strip()):
            p = doc.add_paragraph()
            run = p.add_run(f"[Interactive figure: {line.strip()}]")
            run.font.color.rgb = _MUTED
            run.font.italic = True
            i += 1
            continue

        # Regular paragraph — may contain inline images
        p = doc.add_paragraph()
        segments = _split_inline(line)
        has_inline_image = any(s[0] == "image" for s in segments)

        if has_inline_image:
            for kind, content, extra in segments:
                if kind == "image" and extra:
                    if not _try_add_image(doc, project, extra, width=Inches(4.5)):
                        run = p.add_run(f"[Image: {extra}]")
                        run.font.color.rgb = _MUTED
                else:
                    _add_single_run(p, kind, content, extra)
        else:
            _add_inline_runs(p, line, project)

        i += 1

    # Flush remaining state
    flush_caption()
    flush_table()
    if in_code:
        flush_code()


def _add_single_run(paragraph, kind, content, extra):
    """Add a single styled run to a paragraph."""
    run = paragraph.add_run(content)
    run.font.name = _FONT_SANS
    if kind == "bold":
        run.font.bold = True
    elif kind == "italic":
        run.font.italic = True
    elif kind == "code":
        run.font.name = _FONT_MONO
        run.font.size = Pt(9)
        _shade_run(run, _CODE_BG)
    elif kind == "strike":
        run.font.strike = True
    elif kind == "link":
        run.font.color.rgb = _ACCENT
        run.font.underline = True


def _is_html_embed(text):
    if not text:
        return False
    if re.search(r"\.html?(\?[^\s]*)?$", text, re.I) and re.match(
        r"^(https?://|/api/)", text
    ):
        return True
    m = re.match(r"^!\[([^\]]*)\]\(([^)]+\.html?)(\?[^\s)]*)?\)$", text, re.I)
    return bool(m)
